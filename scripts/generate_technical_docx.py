#!/usr/bin/env python3
"""Render the Uzbek technical documentation Markdown as a styled DOCX file.

Dependency: python-docx. The script intentionally keeps the documentation
content in Markdown so it can be reviewed in version control and regenerated.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BLUE = "2559C7"
DARK = "172033"
MUTED = "48566C"
LIGHT_BLUE = "EDF3FF"
LIGHT_GRAY = "F6F8FB"
BORDER = "CDD6E2"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    properties.append(cannot_split)


def repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_repeat_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), BORDER)


def add_field(run, instruction: str, placeholder: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    run._r.extend([begin, instruction_text, separate, text, end])


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text: str, default_size: float | None = None) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            if default_size:
                run.font.size = Pt(default_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor.from_string(BLUE)
        if default_size and not token.startswith("`"):
            run.font.size = Pt(default_size)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        if default_size:
            run.font.size = Pt(default_size)


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Title": (22, True, BLUE, WD_ALIGN_PARAGRAPH.CENTER),
        "Subtitle": (15, True, MUTED, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 1": (16, True, DARK, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 2": (14, True, BLUE, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": (14, True, DARK, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for name, (size, bold, color, alignment) in heading_specs.items():
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.alignment = alignment
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    styles["Heading 1"].font.all_caps = True
    styles["Heading 1"].paragraph_format.space_before = Pt(0)
    styles["Heading 1"].paragraph_format.space_after = Pt(12)
    styles["Heading 3"].font.italic = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor.from_string(DARK)
    code.paragraph_format.left_indent = Cm(0.5)
    code.paragraph_format.right_indent = Cm(0.5)
    code.paragraph_format.first_line_indent = Cm(0)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.paragraph_format.left_indent = Cm(1.25)
        style.paragraph_format.first_line_indent = Cm(-0.6)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(0)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("SAMARA AI  ·  TEXNIK HUJJAT")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
    bottom = OxmlElement("w:pBdr")
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "4")
    border.set(qn("w:color"), BORDER)
    bottom.append(border)
    paragraph._p.get_or_add_pPr().append(bottom)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    add_field(run, "PAGE", "1")

    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    add_field(run, 'TOC \\o "1-3" \\h \\z \\u', "Mundarijani yangilash uchun Word dasturida Ctrl+A va F9 tugmalarini bosing.")
    note = document.add_paragraph()
    note.paragraph_format.first_line_indent = Cm(0)
    note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.line_spacing = 1.0
    note_run = note.add_run("Izoh: sahifa raqamlari Microsoft Word hujjatni ochganda avtomatik yangilanadi.")
    note_run.italic = True
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = RGBColor.from_string(MUTED)


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph(style="Code Block")
    add_inline(paragraph, "\n".join(lines), 9)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    properties.append(borders)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_borders(table)

    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        prevent_row_split(row)
        if row_index == 0:
            repeat_table_header(row)
        for column_index in range(column_count):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            value = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value, 10.5)
            for run in paragraph.runs:
                run.font.name = "Times New Roman" if run.font.name != "Consolas" else "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw: list[str] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw.append(lines[index].strip())
        index += 1
    rows: list[list[str]] = []
    for row_index, line in enumerate(raw):
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if row_index == 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows, index


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    add_inline(paragraph, text)


def render_markdown(document: Document, source: str) -> None:
    lines = source.splitlines()
    index = 0
    first_heading = True
    on_cover = True
    just_broke = False
    skip_manual_toc = False

    while index < len(lines):
        stripped = lines[index].strip()

        if stripped == "<!-- PAGE BREAK -->":
            document.add_page_break()
            just_broke = True
            on_cover = False
            skip_manual_toc = False
            index += 1
            continue

        if skip_manual_toc:
            if stripped.startswith("#"):
                skip_manual_toc = False
            else:
                index += 1
                continue

        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            add_code_block(document, code_lines)
            just_broke = False
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            just_broke = False
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not first_heading and not just_broke:
                document.add_page_break()
            if on_cover:
                style = "Title" if first_heading else "Subtitle"
            elif text == "MUNDARIJA":
                # Mundarija o'zini TOC ichiga kiritmaslik uchun outline
                # darajasiga ega bo'lmagan uslub qo'llanadi.
                style = "Title"
            else:
                style = f"Heading {level}"
            paragraph = document.add_paragraph(style=style)
            add_inline(paragraph, text)
            if first_heading:
                paragraph.paragraph_format.space_before = Pt(55)
                paragraph.paragraph_format.space_after = Pt(18)
            if text == "MUNDARIJA":
                for run in paragraph.runs:
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor.from_string(DARK)
                add_toc(document)
                skip_manual_toc = True
            first_heading = False
            just_broke = False
            index += 1
            continue

        if stripped == "---":
            paragraph = document.add_paragraph()
            properties = paragraph._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), BORDER)
            borders.append(bottom)
            properties.append(borders)
            index += 1
            continue

        if stripped.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1)
            paragraph.paragraph_format.right_indent = Cm(1)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped[2:].replace("**", ""))
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(BLUE)
            index += 1
            just_broke = False
            continue

        if re.match(r"^-\s+", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, re.sub(r"^-\s+", "", stripped))
            index += 1
            just_broke = False
            continue

        if re.match(r"^\d+\.\s+", stripped):
            # Raqam matnning bir qismi sifatida saqlanadi. Shunda mustaqil
            # bo'limlardagi ro'yxatlar Word tomonidan davom ettirib yuborilmaydi.
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1.25)
            paragraph.paragraph_format.first_line_indent = Cm(-0.6)
            add_inline(paragraph, stripped)
            index += 1
            just_broke = False
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if (
                candidate.startswith("#")
                or candidate.startswith("|")
                or candidate.startswith("```")
                or candidate.startswith("<!-- PAGE BREAK")
                or candidate == "---"
                or candidate.startswith("> ")
                or re.match(r"^-\s+", candidate)
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        add_body_paragraph(document, " ".join(paragraph_lines))
        just_broke = False


def style_cover(document: Document) -> None:
    """Polish paragraphs on the first page without coupling to exact indexes."""
    for paragraph in document.paragraphs:
        if paragraph.text == "HUJJAT PASPORTI":
            break
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.text.startswith(("Tashkilot:", "Oliy ta’lim", "Ish muallifi:", "Ilmiy rahbar:", "Kafedra")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Cm(2)
            paragraph.paragraph_format.space_before = Pt(8)
            for run in paragraph.runs:
                run.font.size = Pt(13)
        if paragraph.text == "Toshkent — 2026":
            paragraph.paragraph_format.space_before = Pt(38)
            for run in paragraph.runs:
                run.bold = True


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: generate_technical_docx.py INPUT.md OUTPUT.docx", file=sys.stderr)
        return 2

    source_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_styles(document)
    configure_page(document)
    render_markdown(document, source_path.read_text(encoding="utf-8"))
    style_cover(document)

    properties = document.core_properties
    properties.title = "Samara AI axborot-tahliliy platformasi — texnik hujjat"
    properties.subject = "Dissertatsiya va diplom ishiga ilova uchun texnik hujjat"
    properties.author = "Samara AI loyiha jamoasi"
    properties.keywords = "Samara AI, EES, ROI, texnik hujjat, sun’iy intellekt"
    properties.comments = "Amaldagi dasturiy kod asosida tayyorlangan."

    document.save(output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
